from pathlib import Path
import pdfplumber  
from docx import Document
from typing import List, Dict, Any
import logging

# note to self - implement image analysis later
# utility importing
from .utils import chunk_text

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file using pdfplumber for better quality
    """
    text = []
    
    try:
        with pdfplumber.open(file_path) as pdf:  # pdfplumber context manager
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract regular text
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                
                # Extract tables and format them
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        # Format table as text with proper spacing
                        table_text = "\n".join(["\t".join(str(cell) for cell in row) for row in table])
                        text.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
                
                logger.debug(f"Processed page {page_num}/{len(pdf.pages)}")
    
    except Exception as e:
        logger.error(f"Error extracting PDF text from {file_path}: {e}")
        raise
    
    return "\n\n".join(text)  # Double newline for page separation


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from DOCX file
    """
    text = []
    
    try:
        doc = Document(file_path)
        
        # Extract both paragraphs and tables
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                para = next((p for p in doc.paragraphs if p._element == element), None)
                if para and para.text.strip():
                    text.append(para.text)
            
            # Check if it's a table
            elif element.tag.endswith('tbl'):
                table = next((t for t in doc.tables if t._element == element), None)
                if table:
                    table_text = []
                    for row in table.rows:
                        row_text = "\t".join(cell.text.strip() for cell in row.cells)
                        table_text.append(row_text)
                    text.append(f"\n[TABLE]\n{chr(10).join(table_text)}\n[/TABLE]\n")
        
    except Exception as e:
        logger.error(f"Error extracting DOCX text from {file_path}: {e}")
        raise
    
    return "\n".join(text)


def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from txt file with fallback encoding"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        logger.warning(f"UTF-8 decode failed for {file_path}, trying latin-1")
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error reading text file {file_path}: {e}")
        raise


def process_document(file_path: Path) -> str:
    """
    Process Document and extract based on file type
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    if not file_path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")
    
    file_type = file_path.suffix.lower()
    
    logger.info(f"Processing {file_type} file: {file_path.name}")
    
    if file_type == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == '.docx':
        return extract_text_from_docx(file_path)
    elif file_type == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported: .pdf, .docx, .txt")


